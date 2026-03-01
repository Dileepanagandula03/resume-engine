from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import sys
import os
import json
from datetime import datetime
import random
import re

# Import your info
from my_info import *
from config import OPENAI_API_KEY

client = OpenAI(api_key=OPENAI_API_KEY)

class Colors:
    GREEN = '\033[92m'
    BLUE = '\033[94m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    END = '\033[0m'
    BOLD = '\033[1m'

def add_hyperlink(paragraph, text, url):
    """Add clickable hyperlink"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)

def add_horizontal_line(paragraph):
    """Add black horizontal line"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def detect_role(jd_text):
    """Detect role from job description"""
    jd_lower = jd_text.lower()
    
    if re.search(r'database engineer|dba|sql developer|database developer', jd_lower):
        return 'Database Engineer'
    elif re.search(r'cloud engineer|aws engineer|azure engineer|cloud architect', jd_lower):
        return 'Cloud Engineer'
    elif re.search(r'data analyst|business analyst|analyst|bi engineer', jd_lower):
        return 'Data Analyst'
    elif re.search(r'machine learning|ml engineer|ai engineer', jd_lower):
        return 'ML Engineer'
    elif re.search(r'data engineer|de|data pipeline', jd_lower):
        return 'Data Engineer'
    else:
        return 'Data Engineer'  # Default

def generate_matched_summary(jd_text):
    """Generate summary with role detection"""
    print(f"{Colors.BLUE}✍️  Generating JD-matched summary...{Colors.END}")
    
    role = detect_role(jd_text)

    prompt = f"""You are an expert resume writer.

JOB DESCRIPTION:
{jd_text}

CANDIDATE BACKGROUND:
{PROFESSIONAL_SUMMARY}

ROLE DETECTED: {role}

YOUR TASK:
Write a 3-4 sentence summary that:
1. Starts with exactly "{role} with 6+ years of experience"
2. Mentions domains: healthcare and financial services
3. Includes key technical strengths relevant to the JD
4. Mentions Master's degree in Information Systems
5. Ends with impact/passion statement
6. Keep it professional and readable

Return ONLY the summary text, no explanations."""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=200
        )
        summary = response.choices[0].message.content.strip()
        summary = summary.replace('**', '').replace('__', '').replace('*', '')
        return summary
    except Exception as e:
        print(f"{Colors.RED}Summary generation failed, using default{Colors.END}")
        return f"{role} with 6+ years of experience architecting data platforms in healthcare and financial services. Proven ability to migrate legacy infrastructures and develop scalable solutions that enhance data accessibility and performance. Recently completed a Master's degree in Information Systems, further enhancing expertise in translating business needs into effective data strategies."

def generate_fallback_bullet(company, slot_num):
    """Generate realistic fallback bullets if AI fails"""
    
    fallbacks = {
        'optum': [
            "Developed Python-based data validation frameworks ensuring healthcare claims data integrity for 5M+ daily records before production loading",
            "Orchestrated cross-functional migration strategies transitioning legacy SQL Server workloads to cloud-native platforms",
            "Engineered automated reconciliation systems comparing source and target data, reducing manual validation effort by 40% across 10+ datasets",
            "Integrated CI/CD pipelines for data services using GitHub Actions, enabling automated testing and deployment with 99.9% success rate",
            "Translated complex business requirements from healthcare stakeholders into scalable technical data solutions adopted by 50+ users",
            "Optimized ETL performance by refactoring PySpark jobs, reducing processing time from 4 hours to 90 minutes for critical claims data"
        ],
        'state_street': [
            "Designed Python/SQL reconciliation frameworks validating 50M+ daily transactions between source systems and enterprise data warehouse",
            "Spearheaded migration of 500GB on-premise data warehouse to Azure Synapse, achieving zero data loss and 40% cost reduction",
            "Constructed ETL pipelines with comprehensive error handling and real-time monitoring, maintaining 99.9% data delivery reliability",
            "Partnered with finance, risk, and operations leaders to resolve complex data discrepancies and improve regulatory reporting accuracy",
            "Standardized data modeling practices across 15+ teams, creating reusable dimension tables that reduced development time by 30%",
            "Implemented infrastructure as code using Terraform, automating cloud resource provisioning and reducing deployment time by 60%"
        ],
        'tech_mahindra': [
            "Formulated SQL transformation logic supporting data migration from Oracle to SQL Server, improving query performance by 60%",
            "Established data validation frameworks that reduced manual investigation time by 30% and ensured 99.5% data reliability for client reporting",
            "Collaborated with international clients to gather requirements and deliver 10+ custom data solutions meeting strict SLA requirements",
            "Streamlined ETL workflows through strategic indexing and parallel processing, cutting batch runtimes from 4 hours to 45 minutes",
            "Generated reusable data quality monitoring dashboards in Tableau, enabling real-time visibility into pipeline health for operations teams",
            "Developed automated alerting systems for pipeline failures, reducing mean time to recovery by 40% for critical production jobs"
        ]
    }
    
    return fallbacks.get(company, fallbacks['tech_mahindra'])[slot_num % len(fallbacks.get(company, fallbacks['tech_mahindra']))]

def extract_jd_technologies(jd_text):
    """Extract ALL technical skills from JD"""
    prompt = f"""Extract EVERY technical skill, tool, technology, and platform from this job description.

JOB DESCRIPTION:
{jd_text}

Return as a comma-separated list. Include:
- Programming languages (Python, SQL, Scala, Java, etc.)
- Tools (Airflow, Spark, Kafka, Tableau, etc.)
- Platforms (Azure, AWS, Snowflake, Databricks, etc.)
- Databases (SQL Server, PostgreSQL, MySQL, Oracle, etc.)
- DevOps tools (Docker, Kubernetes, Terraform, Jenkins, etc.)
- Compliance/security terms (IAM, KMS, HIPAA, GxP, etc.)

ONLY technical items. NO soft skills."""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except:
        return "Python, SQL, Azure, AWS, Spark, Airflow, Databricks, Snowflake, Terraform, Docker, Kubernetes, Tableau"

def get_company_role_name(company, base_role):
    """Get appropriate role name based on company and JD role"""
    role_mapping = {
        'optum': {
            'Database Engineer': 'Database Engineer',
            'Cloud Engineer': 'Cloud Engineer',
            'Data Analyst': 'Data Analyst',
            'ML Engineer': 'ML Engineer',
            'Data Engineer': 'Data Engineer'
        },
        'state_street': {
            'Database Engineer': 'Senior Database Engineer',
            'Cloud Engineer': 'Senior Cloud Engineer',
            'Data Analyst': 'Senior Data Analyst',
            'ML Engineer': 'Senior ML Engineer',
            'Data Engineer': 'Senior Data Engineer'
        },
        'tech_mahindra': {
            'Database Engineer': 'Database Engineer',
            'Cloud Engineer': 'Cloud Engineer',
            'Data Analyst': 'Data Analyst',
            'ML Engineer': 'ML Engineer',
            'Data Engineer': 'Data Engineer'
        }
    }
    return role_mapping.get(company, {}).get(base_role, f"{'Senior ' if company == 'state_street' else ''}{base_role}")

# Core skills groups - ALWAYS present
CORE_SKILLS_GROUPS = [
    "Programming & Scripting: Python, SQL (expert), PySpark, Scala, Bash, PowerShell",
    "Cloud Platforms: Azure (Data Factory, Databricks, Synapse, Functions), AWS (S3, Lambda, Glue, Kinesis)",
    "Data Engineering: ETL/ELT Pipeline Development, Apache Spark, Apache Airflow, Apache Kafka, Data Warehousing",
    "Databases: SQL Server, PostgreSQL, MySQL, CosmosDB, Oracle, Snowflake",
    "DevOps & Tools: Docker, Kubernetes, Jenkins, Git, CI/CD, Terraform",
    "BI & Visualization: Tableau, Power BI, Looker",
    "Compliance & Security: IAM, RBAC, KMS, HIPAA, GxP, Data Governance"
]

def generate_matched_skills(jd_text):
    """Generate skills section - Core groups + JD-specific additions"""
    print(f"{Colors.BLUE}🔧 Generating JD-matched skills...{Colors.END}")
    
    # Extract all tech from JD
    jd_tech_skills = extract_jd_technologies(jd_text)
    jd_tech_list = [t.strip() for t in jd_tech_skills.split(',')]
    
    # Start with core skills (these NEVER change)
    base_skills = CORE_SKILLS_GROUPS.copy()
    
    # Find JD skills NOT already in core
    missing_skills = []
    for tech in jd_tech_list:
        tech_lower = tech.lower()
        found = False
        
        # Check if tech already exists in core skills
        for core_group in base_skills:
            if tech_lower in core_group.lower():
                found = True
                break
        
        if not found and len(tech) > 2:  # Avoid single chars
            missing_skills.append(tech)
    
    # Add missing skills as additional group
    if missing_skills:
        # Remove duplicates while preserving order
        unique_missing = []
        for tech in missing_skills:
            if tech not in unique_missing:
                unique_missing.append(tech)
        
        extra_group = "Additional Technologies: " + ", ".join(unique_missing[:8])  # Limit to 8
        base_skills.append(extra_group)
    
    print(f"{Colors.GREEN}✅ Core skills preserved + {len(missing_skills)} JD-specific skills added{Colors.END}")
    return base_skills

def generate_matching_bullets(jd_text):
    """Generate bullets with tech mapping based on JD"""
    print(f"\n{Colors.BLUE}🤖 AI analyzing JD and generating bullets...{Colors.END}")

    # Get JD technologies
    jd_tech = extract_jd_technologies(jd_text)
    role = detect_role(jd_text)

    prompt = f"""You are an expert resume writer. Your job is to create bullets that SHOW real experience with EVERY technology from the JD.

JOB DESCRIPTION TECHNOLOGIES (must include these in bullets):
{jd_tech}

DETECTED ROLE: {role}

CANDIDATE BACKGROUND:
{YOUR_EXPERIENCE}

TECHNOLOGY TO COMPANY MAPPING GUIDELINES:
- Snowflake, Terraform, Airflow, Jenkins, Docker, Kubernetes → Associate with State Street (senior financial role)
- Databricks, Kafka, Azure Functions, Event Hubs → Associate with Optum (current healthcare role)
- Tableau, SSIS, client delivery → Associate with Tech Mahindra (early career telecom role)
- Python, SQL → Universal (use across all)

YOUR TASK:
Create 6 bullets for EACH company below.
Each bullet MUST:
1. Include specific technologies from the JD mapped to appropriate company
2. Show real production experience with metrics (%, millions of records, etc.)
3. Sound like actual work done at that specific company
4. Be 20-30 words
5. Start with strong action verb
6. All 18 bullets must have UNIQUE action verbs (no repeats)

WRITE BULLETS FOR:

OPTUM (Healthcare - Current Role):
- Focus on Python, data validation, healthcare data, production pipelines
- Include metrics (5M+ records, 99.8% accuracy, 30-40% improvements)
- Use Databricks, Kafka, Azure Functions where JD mentions them

STATE STREET (Financial - Senior Role):
- Focus on large-scale data, migrations, SQL, Spark, Airflow
- Include metrics (50M+ records, 500GB data, 40% cost reduction)
- Use Snowflake, Terraform, Jenkins, Docker, K8s where JD mentions them

TECH MAHINDRA (Telecom - Early Role):
- Focus on automation, data quality, SQL, client delivery
- Include metrics (2M+ daily records, 50% reduction, 60% faster)
- Use Tableau, SSIS, client delivery where JD mentions them

OUTPUT FORMAT:

OPTUM:
1. [bullet with Python + validation + metric]
2. [bullet with Databricks/Spark + pipeline + metric]
3. [bullet with SQL + data validation + metric]
4. [bullet with cloud migration + Azure + metric]
5. [bullet with production support + monitoring + metric]
6. [bullet with data quality/Kafka + streaming + metric]

STATE STREET:
1. [bullet with Snowflake/Spark + large-scale processing + metric]
2. [bullet with Airflow + pipeline orchestration + metric]
3. [bullet with Terraform/IaC + cloud migration + metric]
4. [bullet with data modeling + warehousing + metric]
5. [bullet with stakeholder collaboration + requirements + metric]
6. [bullet with Jenkins/Docker + CI/CD + metric]

TECH MAHINDRA:
1. [bullet with Python automation + data quality + metric]
2. [bullet with SQL optimization + performance + metric]
3. [bullet with SSIS/ETL + pipeline + metric]
4. [bullet with client delivery + requirements + metric]
5. [bullet with Tableau/BI + efficiency + metric]
6. [bullet with monitoring/alerting + troubleshooting + metric]

Remember: 18 bullets total, 18 UNIQUE action verbs. No repeats."""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """You are a senior data engineer writing real resume bullets.
                    CRITICAL RULE: Every single bullet across ALL companies must start with a DIFFERENT action verb.
                    18 bullets = 18 unique verbs. No repeats.
                    Match technologies to appropriate companies based on seniority.
                    Include concrete numbers and metrics. Make it sound like real work done."""
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=3500
        )

        bullets_text = response.choices[0].message.content
        bullets = parse_bullets(bullets_text)

        # Ensure we have 6 bullets per company
        for company in bullets:
            while len(bullets[company]) < 6:
                bullets[company].append(generate_fallback_bullet(company, len(bullets[company])))
            bullets[company] = bullets[company][:6]

        print(f"{Colors.GREEN}✅ Generated bullets with JD technologies{Colors.END}")
        return bullets

    except Exception as e:
        print(f"{Colors.RED}Error generating bullets, using fallback{Colors.END}")
        return {
            'optum': [generate_fallback_bullet('optum', i) for i in range(6)],
            'state_street': [generate_fallback_bullet('state_street', i) for i in range(6)],
            'tech_mahindra': [generate_fallback_bullet('tech_mahindra', i) for i in range(6)]
        }

def parse_bullets(text):
    bullets = {'optum': [], 'state_street': [], 'tech_mahindra': []}
    lines = text.strip().split('\n')
    current_company = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        upper_line = line.upper()
        if 'OPTUM' in upper_line:
            current_company = 'optum'
            continue
        elif 'STATE STREET' in upper_line:
            current_company = 'state_street'
            continue
        elif 'TECH MAHINDRA' in upper_line:
            current_company = 'tech_mahindra'
            continue
        
        if current_company and line and (line[0].isdigit() or line[0] in ['-', '•', '·']):
            bullet = line.lstrip('0123456789.-•·) ').strip()
            if len(bullet) > 15 and 'bullet with' not in bullet.lower():
                bullets[current_company].append(bullet)
    
    return bullets

def calculate_ats_score(resume_text, jd_text):
    print(f"\n{Colors.BLUE}📊 Calculating ATS score...{Colors.END}")
    
    prompt = f"""Extract ALL technical skills and keywords from this JD.

{jd_text}

Return comma-separated list. Include ONLY technical items:
- Programming languages
- Tools and technologies
- Platforms
- Databases
- Technical methodologies
- Compliance/security terms

NO soft skills."""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=300
        )
        
        jd_keywords = [k.strip() for k in response.choices[0].message.content.split(',')]
        resume_lower = resume_text.lower()
        matched = [kw for kw in jd_keywords if kw.lower() in resume_lower]
        
        score = (len(matched) / len(jd_keywords)) * 100 if jd_keywords else 0
        
        print(f"\n{Colors.GREEN}{Colors.BOLD}✅ ATS SCORE: {score:.1f}%{Colors.END}")
        print(f"\n{Colors.GREEN}✅ Matched ({len(matched)}): {', '.join(matched[:15])}...{Colors.END}")
        
        missing = [kw for kw in jd_keywords if kw not in matched]
        if missing:
            print(f"\n{Colors.YELLOW}⚠️ Missing ({len(missing)}): {', '.join(missing[:10])}...{Colors.END}")
        
        return score
    except Exception as e:
        print(f"{Colors.RED}ATS scoring failed: {e}{Colors.END}")
        return 0

def generate_docx(bullets, job_title, matched_summary, jd_text):
    print(f"\n{Colors.BLUE}📄 Generating 1-page resume...{Colors.END}")
    
    doc = Document()
    
    # TIGHTER MARGINS (0.3")
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
    
    # Set default paragraph spacing to ZERO
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.0
    
    # NAME
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(YOUR_NAME)
    name_run.font.size = Pt(14)
    name_run.bold = True
    
    # CONTACT
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)
    contact_run = contact_para.add_run(f"{YOUR_PHONE} | {YOUR_EMAIL} | {YOUR_LOCATION} | ")
    contact_run.font.size = Pt(9)
    
    add_hyperlink(contact_para, 'LinkedIn', YOUR_LINKEDIN)
    contact_para.add_run(' | ').font.size = Pt(9)
    add_hyperlink(contact_para, 'GitHub', YOUR_GITHUB)
    contact_para.add_run(' | ').font.size = Pt(9)
    add_hyperlink(contact_para, 'Portfolio', YOUR_PORTFOLIO)
    
    # SUMMARY - Header above line
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_heading.paragraph_format.space_before = Pt(4)
    summary_heading.paragraph_format.space_after = Pt(0)
    summary_run = summary_heading.add_run('Summary')
    summary_run.font.size = Pt(11)
    summary_run.bold = True
    summary_run.font.color.rgb = RGBColor(0, 86, 193)
    
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    summary_para = doc.add_paragraph(matched_summary)
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(4)
    for run in summary_para.runs:
        run.font.size = Pt(10)
    
    # SKILLS - Header above line
    skills_heading = doc.add_paragraph()
    skills_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    skills_heading.paragraph_format.space_before = Pt(0)
    skills_heading.paragraph_format.space_after = Pt(0)
    skills_run = skills_heading.add_run('Skills')
    skills_run.font.size = Pt(11)
    skills_run.bold = True
    skills_run.font.color.rgb = RGBColor(0, 86, 193)
    
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Generate skills (Core + JD-specific)
    matched_skills = generate_matched_skills(jd_text)
    
    for skill_line in matched_skills:
        if ':' in skill_line:
            category, skills = skill_line.split(':', 1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            cat_run = p.add_run(category + ':')
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            skill_run = p.add_run(skills)
            skill_run.font.size = Pt(10)
    
    # EXPERIENCE - Header above line
    exp_heading = doc.add_paragraph()
    exp_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exp_heading.paragraph_format.space_before = Pt(4)
    exp_heading.paragraph_format.space_after = Pt(0)
    exp_run = exp_heading.add_run('Experience')
    exp_run.font.size = Pt(11)
    exp_run.bold = True
    exp_run.font.color.rgb = RGBColor(0, 86, 193)
    
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Detect role for company titles
    base_role = detect_role(jd_text)
    
    # OPTUM
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]
    
    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    optum_role = get_company_role_name('optum', base_role)
    comp_run = left_para.add_run(f'{optum_role}, Optum')
    comp_run.bold = True
    comp_run.font.size = Pt(10)
    left_para.add_run(' | Chicago, IL').font.size = Pt(10)
    
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    date_run = right_para.add_run('01/2025 - Present')
    date_run.font.size = Pt(10)
    date_run.bold = True
    
    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    for bullet in bullets['optum'][:6]:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(0)
    
    # STATE STREET
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]
    
    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    ss_role = get_company_role_name('state_street', base_role)
    comp_run = left_para.add_run(f'{ss_role}, State Street Corporation')
    comp_run.bold = True
    comp_run.font.size = Pt(10)
    left_para.add_run(' | Boston, MA').font.size = Pt(10)
    
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    date_run = right_para.add_run('05/2021 - 12/2023')
    date_run.font.size = Pt(10)
    date_run.bold = True
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    for bullet in bullets['state_street'][:6]:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(0)
    
    # TECH MAHINDRA
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]
    
    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    tm_role = get_company_role_name('tech_mahindra', base_role)
    comp_run = left_para.add_run(f'{tm_role}, Tech Mahindra')
    comp_run.bold = True
    comp_run.font.size = Pt(10)
    left_para.add_run(' | Hyderabad, India').font.size = Pt(10)
    
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    date_run = right_para.add_run('09/2018 - 04/2021')
    date_run.font.size = Pt(10)
    date_run.bold = True
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    for bullet in bullets['tech_mahindra'][:6]:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # PROJECTS SECTION
    proj_heading = doc.add_paragraph()
    proj_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_heading.paragraph_format.space_before = Pt(4)
    proj_heading.paragraph_format.space_after = Pt(0)
    proj_run = proj_heading.add_run('Projects')
    proj_run.font.size = Pt(11)
    proj_run.bold = True
    proj_run.font.color.rgb = RGBColor(0, 86, 193)
    
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Project 1
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1_title = p1.add_run('Azure Medallion Data Pipeline')
    p1_title.bold = True
    p1_title.font.size = Pt(10)
    p1.add_run(' ').font.size = Pt(10)
    add_hyperlink(p1, 'View Project', 'https://github.com/Dileepanagandula03/azure-medallion-data-pipeline')
    
    p1_desc = doc.add_paragraph('Built a Medallion architecture leveraging ADF, ADLS, and Databricks to orchestrate ETL workflows, apply data quality checks, and deliver curated datasets consumed by Power BI dashboards.')
    p1_desc.paragraph_format.space_before = Pt(0)
    p1_desc.paragraph_format.space_after = Pt(4)  # Slightly more space between projects
    p1_desc.paragraph_format.left_indent = Inches(0.25)
    for run in p1_desc.runs:
        run.font.size = Pt(9)
    
    # Project 2
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2_title = p2.add_run('AI-Powered Resume Tailoring Tool')
    p2_title.bold = True
    p2_title.font.size = Pt(10)
    p2.add_run(' | Recently developed').font.size = Pt(9)
    p2.add_run(' ').font.size = Pt(10)
    add_hyperlink(p2, 'View Project', 'https://github.com/Dileepanagandula03/azure-medallion-data-pipeline')
    
    p2_desc = doc.add_paragraph('Built an AI-powered resume tailoring tool using OpenAI\'s LLM APIs and Python that automates resume customization in under 2 minutes, improves ATS match scores by 25-40%, and delivers 75% cost savings compared to commercial solutions.')
    p2_desc.paragraph_format.space_before = Pt(0)
    p2_desc.paragraph_format.space_after = Pt(4)
    p2_desc.paragraph_format.left_indent = Inches(0.25)
    for run in p2_desc.runs:
        run.font.size = Pt(9)
    
    # Project 3
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    p3_title = p3.add_run('Real-Time and Batch Weather Data Pipeline')
    p3_title.bold = True
    p3_title.font.size = Pt(10)
    p3.add_run(' ').font.size = Pt(10)
    add_hyperlink(p3, 'View Project', 'https://github.com/Dileepanagandula03/weather-pipeline-project')
    
    p3_desc = doc.add_paragraph('Built an end-to-end real-time and batch weather data pipeline on Azure, leveraging serverless ingestion and streaming analytics to process API-based event data, apply transformations and validations, and deliver analytics-ready datasets with cost-optimized architecture.')
    p3_desc.paragraph_format.space_before = Pt(0)
    p3_desc.paragraph_format.space_after = Pt(4)
    p3_desc.paragraph_format.left_indent = Inches(0.25)
    for run in p3_desc.runs:
        run.font.size = Pt(9)
    
    # Project 4
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)
    p4_title = p4.add_run('ETL vs ELT Decision Framework Capstone Project')
    p4_title.bold = True
    p4_title.font.size = Pt(10)
    p4.add_run(' ').font.size = Pt(10)
    add_hyperlink(p4, 'View Paper', 'https://github.com/Dileepanagandula03/Capstone-Data-Engineering-Pipelines')
    
    p4_desc = doc.add_paragraph('Developed a comparative ETL vs ELT decision framework by evaluating performance, scalability, cost efficiency, and governance tradeoffs in cloud-based data pipelines, providing actionable guidance for selecting architectures.')
    p4_desc.paragraph_format.space_before = Pt(0)
    p4_desc.paragraph_format.space_after = Pt(4)
    p4_desc.paragraph_format.left_indent = Inches(0.25)
    for run in p4_desc.runs:
        run.font.size = Pt(9)

    # EDUCATION SECTION
    edu_heading = doc.add_paragraph()
    edu_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edu_heading.paragraph_format.space_before = Pt(4)
    edu_heading.paragraph_format.space_after = Pt(0)
    edu_run = edu_heading.add_run('Education')
    edu_run.font.size = Pt(11)
    edu_run.bold = True
    edu_run.font.color.rgb = RGBColor(0, 86, 193)
    
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Education line 1
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]
    
    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.add_run('Webster University, Master of Science in Information Systems').font.size = Pt(10)
    
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.add_run('Jan 2024 – Dec 2025').font.size = Pt(10)
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # Education line 2
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]
    
    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.add_run('St. Louis, MO').font.size = Pt(10)
    
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.add_run('3.8/4.0').font.size = Pt(10)
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # Save
    if not os.path.exists('output'):
        os.makedirs('output')
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    job_title_clean = job_title.replace(' ', '_').replace('/', '_')
    
    word_filename = f"output/{job_title_clean}_{timestamp}.docx"
    doc.save(word_filename)
    print(f"{Colors.GREEN}✅ Word document created: {word_filename}{Colors.END}")
    
    pdf_filename = f"output/{job_title_clean}_{timestamp}.pdf"
    try:
        print(f"{Colors.BLUE}🔄 Converting to PDF...{Colors.END}")
        convert(word_filename, pdf_filename)
        print(f"{Colors.GREEN}✅ PDF saved: {pdf_filename}{Colors.END}")
        os.remove(word_filename)
        return pdf_filename, doc
    except Exception as e:
        print(f"{Colors.RED}⚠️ PDF conversion failed: {e}{Colors.END}")
        return word_filename, doc

def main():
    print(f"{Colors.BOLD}{'='*60}{Colors.END}")
    print(f"{Colors.BOLD}{Colors.BLUE}   JD MATCHING RESUME GENERATOR{Colors.END}")
    print(f"{Colors.BOLD}{'='*60}{Colors.END}")
    
    print(f"\n{Colors.YELLOW}Paste job description below.{Colors.END}")
    print(f"{Colors.YELLOW}Press Ctrl+Z then Enter when done:{Colors.END}\n")
    
    jd_lines = []
    try:
        while True:
            line = input()
            jd_lines.append(line)
    except EOFError:
        pass
    
    jd_text = '\n'.join(jd_lines)
    
    if not jd_text.strip():
        print(f"{Colors.RED}❌ No JD provided.{Colors.END}")
        return
    
    job_title = input(f"\n{Colors.YELLOW}Enter job title: {Colors.END}").strip()
    if not job_title:
        job_title = "Resume"
    
    # Detect role
    role = detect_role(jd_text)
    print(f"{Colors.BLUE}📋 Detected Role: {role}{Colors.END}")
    
    # Generate matched summary with role detection
    matched_summary = generate_matched_summary(jd_text)
    print(f"\n{Colors.GREEN}✅ Summary generated{Colors.END}")
    
    # Generate bullets with tech mapping
    bullets = generate_matching_bullets(jd_text)
    
    # Display generated bullets
    print(f"\n{Colors.BOLD}{'='*60}{Colors.END}")
    print(f"{Colors.BOLD}GENERATED BULLETS:{Colors.END}")
    print(f"{Colors.BOLD}{'='*60}{Colors.END}")
    
    for company, company_bullets in bullets.items():
        if company_bullets:
            company_display = {
                'optum': 'OPTUM',
                'state_street': 'STATE STREET',
                'tech_mahindra': 'TECH MAHINDRA'
            }.get(company, company.upper())
            print(f"\n{Colors.BLUE}{company_display}:{Colors.END}")
            for i, bullet in enumerate(company_bullets[:6], 1):
                print(f"{i}. {bullet}")
    
    # Generate docx
    filename, doc = generate_docx(bullets, job_title, matched_summary, jd_text)
    
    # ATS score
    resume_text = '\n'.join([p.text for p in doc.paragraphs])
    ats_score = calculate_ats_score(resume_text, jd_text)
    
    print(f"\n{Colors.BOLD}{'='*60}{Colors.END}")
    print(f"{Colors.GREEN}{Colors.BOLD}🎉 DONE! Check your output folder!{Colors.END}")
    print(f"{Colors.BOLD}{'='*60}{Colors.END}\n")

if __name__ == "__main__":
    main()
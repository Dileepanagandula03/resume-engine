# templates/experience_template.py

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.docx_utils import add_horizontal_line
import user_session


def _role_for_company(exp, detected_role):
    """Adapt job title: preserve seniority prefix from user's actual role, swap role noun."""
    user_role = exp["role"]
    seniority_prefixes = ["senior", "lead", "principal", "staff", "junior", "associate"]
    prefix = ""
    for w in user_role.lower().split():
        if w in seniority_prefixes:
            prefix = w.capitalize() + " "
            break
    return f"{prefix}{detected_role}"


def _remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_experience_section(doc, bullets, detected_role):
    """Add the experience section using companies from my_info.YOUR_EXPERIENCE."""

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("Experience")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 86, 193)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line)

    for idx, exp in enumerate(user_session.YOUR_EXPERIENCE):
        if idx > 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(2)
            spacer.paragraph_format.space_after = Pt(0)

        company_bullets = bullets.get(exp["key"], [])
        display_role = _role_for_company(exp, detected_role)
        _add_company_block(doc, exp, display_role, company_bullets)


def _add_company_block(doc, exp, display_role, company_bullets):
    """Add one company block: title row + bullet points."""

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]

    left = row.cells[0].paragraphs[0]
    left.paragraph_format.space_before = Pt(0)
    left.paragraph_format.space_after = Pt(0)
    title_run = left.add_run(f"{display_role}, {exp['company']}")
    title_run.bold = True
    title_run.font.size = Pt(10)
    if exp.get("location"):
        left.add_run(f" | {exp['location']}").font.size = Pt(10)

    right = row.cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_before = Pt(0)
    right.paragraph_format.space_after = Pt(0)
    date_str = f"{exp['start_date']} \u2013 {exp['end_date']}"
    date_run = right.add_run(date_str)
    date_run.font.size = Pt(10)
    date_run.bold = True

    _remove_table_borders(table)

    for bullet in company_bullets[:6]:
        p = doc.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)

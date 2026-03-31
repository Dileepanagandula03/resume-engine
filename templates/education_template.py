# templates/education_template.py

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.docx_utils import add_horizontal_line
import user_session


def _remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_education_section(doc):
    """Add education section from my_info.YOUR_EDUCATION."""

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("Education")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 86, 193)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line)

    for edu in user_session.YOUR_EDUCATION:
        _add_education_entry(doc, edu)


def _add_education_entry(doc, edu):
    """Add one education entry (school + degree row, then location + GPA row)."""

    # Row 1: school name + degree — date
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]

    left = row.cells[0].paragraphs[0]
    left.paragraph_format.space_before = Pt(0)
    left.paragraph_format.space_after = Pt(0)
    left.add_run(f"{edu['school']}, {edu['degree']}").font.size = Pt(10)

    right = row.cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_before = Pt(0)
    right.paragraph_format.space_after = Pt(0)
    start = edu.get("start_date", "")
    end = edu.get("end_date", "")
    if start and end:
        date_str = f"{start} \u2013 {end}"
    elif start or end:
        date_str = start or end
    else:
        date_str = ""
    if date_str:
        right.add_run(date_str).font.size = Pt(10)

    _remove_table_borders(table)

    # Row 2: location — GPA (only if either is present)
    location = edu.get("location", "")
    gpa = edu.get("gpa", "")
    if location or gpa:
        table2 = doc.add_table(rows=1, cols=2)
        row2 = table2.rows[0]

        left2 = row2.cells[0].paragraphs[0]
        left2.paragraph_format.space_before = Pt(0)
        left2.paragraph_format.space_after = Pt(0)
        left2.add_run(location).font.size = Pt(10)

        right2 = row2.cells[1].paragraphs[0]
        right2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right2.paragraph_format.space_before = Pt(0)
        right2.paragraph_format.space_after = Pt(0)
        if gpa:
            right2.add_run(gpa).font.size = Pt(10)

        _remove_table_borders(table2)

# templates/skills_template.py

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_utils import add_horizontal_line

def add_skills_section(doc, skills_list):
    """Add skills section with core + JD-specific skills"""
    
    # Header
    skills_heading = doc.add_paragraph()
    skills_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    skills_heading.paragraph_format.space_before = Pt(0)
    skills_heading.paragraph_format.space_after = Pt(0)
    skills_run = skills_heading.add_run('Skills')
    skills_run.font.size = Pt(11)
    skills_run.bold = True
    skills_run.font.color.rgb = RGBColor(0, 86, 193)
    
    # Line
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Skills
    for skill_line in skills_list:
        if ':' in skill_line:
            category, skills = skill_line.split(':', 1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            cat_run = p.add_run(category + ':')
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            skill_run = p.add_run(skills)
            skill_run.font.size = Pt(10)
# templates/summary_template.py

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_utils import add_horizontal_line

def add_summary_section(doc, summary_text):
    """Add summary section with header above line"""
    
    # Header
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_heading.paragraph_format.space_before = Pt(4)
    summary_heading.paragraph_format.space_after = Pt(0)
    summary_run = summary_heading.add_run('Summary')
    summary_run.font.size = Pt(11)
    summary_run.bold = True
    summary_run.font.color.rgb = RGBColor(0, 86, 193)
    
    # Line
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line_para)
    
    # Summary text
    summary_para = doc.add_paragraph(summary_text)
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(4)
    for run in summary_para.runs:
        run.font.size = Pt(10)
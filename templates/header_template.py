# templates/header_template.py

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_utils import add_hyperlink

def add_header(doc, name, phone, email, location, linkedin, github, portfolio):
    """Add name and contact information"""
    
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(name)
    name_run.font.size = Pt(14)
    name_run.bold = True
    
    # Contact
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)
    contact_run = contact_para.add_run(f"{phone} | {email} | {location} | ")
    contact_run.font.size = Pt(9)
    
    add_hyperlink(contact_para, 'LinkedIn', linkedin)
    contact_para.add_run(' | ').font.size = Pt(9)
    add_hyperlink(contact_para, 'GitHub', github)
    contact_para.add_run(' | ').font.size = Pt(9)
    add_hyperlink(contact_para, 'Portfolio', portfolio)
# templates/projects_template.py

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_utils import add_horizontal_line, add_hyperlink
import user_session


def add_projects_section(doc):
    """Add projects section from my_info.YOUR_PROJECTS."""

    if not user_session.YOUR_PROJECTS:
        return  # Skip entire section when no projects were parsed

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run("Projects")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 86, 193)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line)

    for project in user_session.YOUR_PROJECTS:
        _add_project(
            doc,
            title=project["title"],
            url=project.get("url", ""),
            description=project["description"],
            link_text=project.get("link_text", "View Project"),
        )


def _add_project(doc, title, url, description, link_text="View Project"):
    """Add a single project entry."""

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10)

    if url:
        p.add_run("  ").font.size = Pt(10)
        add_hyperlink(p, link_text, url)

    desc = doc.add_paragraph(description)
    desc.paragraph_format.space_before = Pt(0)
    desc.paragraph_format.space_after = Pt(4)
    desc.paragraph_format.left_indent = Inches(0.25)
    for run in desc.runs:
        run.font.size = Pt(9)
